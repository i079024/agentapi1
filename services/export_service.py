"""
Export Service - JSON and Word document export functionality
Handles test and result export in multiple formats
"""

import json
import os
import tempfile
from typing import Dict, List, Any, Optional
from datetime import datetime
from io import BytesIO

# Try to import docx for Word export, fallback to text if not available
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ExportService:
    def __init__(self):
        self.export_formats = ["json", "word", "html", "csv"]
        
    def export_tests(self, tests: List[Dict[str, Any]], format_type: str = "json") -> Dict[str, Any]:
        """Export tests in specified format"""
        if format_type not in self.export_formats:
            raise ValueError(f"Unsupported format: {format_type}")
        
        if format_type == "json":
            return self._export_tests_json(tests)
        elif format_type == "word":
            return self._export_tests_word(tests)
        elif format_type == "html":
            return self._export_tests_html(tests)
        elif format_type == "csv":
            return self._export_tests_csv(tests)
        else:
            return {"success": False, "error": f"Unsupported format: {format_type}"}
    
    def export_results(self, results: Dict[str, Any], format_type: str = "json") -> Dict[str, Any]:
        """Export test results in specified format"""
        if format_type not in self.export_formats:
            raise ValueError(f"Unsupported format: {format_type}")
        
        if format_type == "json":
            return self._export_results_json(results)
        elif format_type == "word":
            return self._export_results_word(results)
        elif format_type == "html":
            return self._export_results_html(results)
        elif format_type == "csv":
            return self._export_results_csv(results)
        else:
            return {"success": False, "error": f"Unsupported format: {format_type}"}
    
    def _export_tests_json(self, tests: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Export tests as JSON"""
        export_data = {
            "export_metadata": {
                "type": "tests",
                "format": "json",
                "created_at": datetime.now().isoformat(),
                "version": "1.0",
                "total_tests": len(tests),
                "tool": "Agent API Testing Platform"
            },
            "tests": tests
        }
        
        # Create JSON string
        json_content = json.dumps(export_data, indent=2, default=str)
        
        return {
            "success": True,
            "format": "json",
            "content": json_content,
            "filename": f"api_tests_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            "size": len(json_content.encode('utf-8'))
        }
    
    def _export_results_json(self, results: Dict[str, Any]) -> Dict[str, Any]:
        """Export test results as JSON"""
        export_data = {
            "export_metadata": {
                "type": "results",
                "format": "json",
                "created_at": datetime.now().isoformat(),
                "version": "1.0",
                "tool": "Agent API Testing Platform"
            },
            "results": results
        }
        
        # Create JSON string
        json_content = json.dumps(export_data, indent=2, default=str)
        
        return {
            "success": True,
            "format": "json",
            "content": json_content,
            "filename": f"test_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            "size": len(json_content.encode('utf-8'))
        }
    
    def _export_tests_word(self, tests: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Export tests as Word document"""
        if not DOCX_AVAILABLE:
            return self._export_fallback_text(tests, "tests")
        
        try:
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading('API Test Suite', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_heading('Export Information', level=1)
            metadata_table = doc.add_table(rows=4, cols=2)
            metadata_table.style = 'Table Grid'
            
            metadata_rows = [
                ['Export Date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
                ['Total Tests', str(len(tests))],
                ['Tool', 'Agent API Testing Platform'],
                ['Format Version', '1.0']
            ]
            
            for i, (key, value) in enumerate(metadata_rows):
                metadata_table.cell(i, 0).text = key
                metadata_table.cell(i, 1).text = value
            
            # Add test details
            doc.add_heading('Test Details', level=1)
            
            for i, test in enumerate(tests, 1):
                # Test header
                doc.add_heading(f'Test {i}: {test.get("name", "Unnamed Test")}', level=2)
                
                # Test details table
                test_table = doc.add_table(rows=8, cols=2)
                test_table.style = 'Table Grid'
                
                test_details = [
                    ['Method', test.get('method', 'GET')],
                    ['URL', test.get('url', '')],
                    ['Description', test.get('description', 'No description')],
                    ['Headers', json.dumps(test.get('headers', {}), indent=2)],
                    ['Body', json.dumps(test.get('body'), indent=2) if test.get('body') else 'None'],
                    ['Timeout', f"{test.get('timeout', 30)} seconds"],
                    ['Tags', ', '.join(test.get('tags', []))],
                    ['Assertions', f"{len(test.get('assertions', []))} configured"]
                ]
                
                for j, (key, value) in enumerate(test_details):
                    test_table.cell(j, 0).text = key
                    test_table.cell(j, 1).text = str(value)
                
                # Add assertions details
                assertions = test.get('assertions', [])
                if assertions:
                    doc.add_heading('Assertions', level=3)
                    for k, assertion in enumerate(assertions, 1):
                        p = doc.add_paragraph(f"{k}. {assertion.get('type', 'unknown')} - {assertion.get('description', 'No description')}")
            
            # Save to BytesIO
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return {
                "success": True,
                "format": "word",
                "content": doc_buffer.getvalue(),
                "filename": f"api_tests_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "size": len(doc_buffer.getvalue()),
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Failed to create Word document: {str(e)}",
                "fallback": self._export_fallback_text(tests, "tests")
            }
    
    def _export_results_word(self, results: Dict[str, Any]) -> Dict[str, Any]:
        """Export test results as Word document"""
        if not DOCX_AVAILABLE:
            return self._export_fallback_text([results], "results")
        
        try:
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading('API Test Results Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add executive summary
            doc.add_heading('Executive Summary', level=1)
            
            summary = results.get('summary', {})
            if summary:
                summary_table = doc.add_table(rows=4, cols=2)
                summary_table.style = 'Table Grid'
                
                summary_rows = [
                    ['Total Tests', str(summary.get('total_tests', 0))],
                    ['Passed', str(summary.get('passed', 0))],
                    ['Failed', str(summary.get('failed', 0))],
                    ['Success Rate', f"{summary.get('success_rate', 0)}%"]
                ]
                
                for i, (key, value) in enumerate(summary_rows):
                    summary_table.cell(i, 0).text = key
                    summary_table.cell(i, 1).text = value
            
            # Add performance metrics
            execution_report = results.get('execution_report', {})
            performance = execution_report.get('performance_metrics', {})
            
            if performance:
                doc.add_heading('Performance Metrics', level=1)
                perf_table = doc.add_table(rows=4, cols=2)
                perf_table.style = 'Table Grid'
                
                perf_rows = [
                    ['Total Execution Time', f"{performance.get('total_execution_time', 0):.2f}s"],
                    ['Average Response Time', f"{performance.get('average_response_time', 0):.2f}s"],
                    ['Fastest Test', f"{performance.get('fastest_test', 0):.2f}s"],
                    ['Slowest Test', f"{performance.get('slowest_test', 0):.2f}s"]
                ]
                
                for i, (key, value) in enumerate(perf_rows):
                    perf_table.cell(i, 0).text = key
                    perf_table.cell(i, 1).text = value
            
            # Add detailed test results
            test_results = results.get('test_results', [])
            if test_results:
                doc.add_heading('Detailed Test Results', level=1)
                
                for i, test_result in enumerate(test_results, 1):
                    test = test_result.get('test', {})
                    success = test_result.get('success', False)
                    
                    # Test result header
                    status = "✅ PASSED" if success else "❌ FAILED"
                    doc.add_heading(f'Test {i}: {test.get("name", "Unnamed")} - {status}', level=2)
                    
                    # Test result details
                    result_table = doc.add_table(rows=6, cols=2)
                    result_table.style = 'Table Grid'
                    
                    response = test_result.get('response', {})
                    result_details = [
                        ['Status Code', str(response.get('status_code', 'N/A'))],
                        ['Response Time', f"{test_result.get('execution_time', 0):.3f}s"],
                        ['Assertions Passed', str(test_result.get('assertions_passed', 0))],
                        ['Assertions Failed', str(test_result.get('assertions_failed', 0))],
                        ['Error Message', test_result.get('error_message', 'None')],
                        ['Request URL', f"{test.get('method', 'GET')} {test.get('url', '')}"]
                    ]
                    
                    for j, (key, value) in enumerate(result_details):
                        result_table.cell(j, 0).text = key
                        result_table.cell(j, 1).text = str(value)
            
            # Add failure analysis if available
            failure_analysis = execution_report.get('failure_analysis')
            if failure_analysis and failure_analysis.get('failed_tests'):
                doc.add_heading('Failure Analysis', level=1)
                doc.add_paragraph('Failed Tests:')
                for failed_test in failure_analysis.get('failed_tests', []):
                    doc.add_paragraph(f"• {failed_test}", style='List Bullet')
                
                if failure_analysis.get('common_errors'):
                    doc.add_paragraph('Common Errors:')
                    for error in failure_analysis.get('common_errors', []):
                        doc.add_paragraph(f"• {error}", style='List Bullet')
            
            # Save to BytesIO
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return {
                "success": True,
                "format": "word",
                "content": doc_buffer.getvalue(),
                "filename": f"test_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "size": len(doc_buffer.getvalue()),
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Failed to create Word document: {str(e)}",
                "fallback": self._export_fallback_text([results], "results")
            }
    
    def _export_tests_html(self, tests: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Export tests as HTML"""
        html_content = self._generate_tests_html(tests)
        
        return {
            "success": True,
            "format": "html",
            "content": html_content,
            "filename": f"api_tests_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
            "size": len(html_content.encode('utf-8')),
            "content_type": "text/html"
        }
    
    def _export_results_html(self, results: Dict[str, Any]) -> Dict[str, Any]:
        """Export results as HTML"""
        html_content = self._generate_results_html(results)
        
        return {
            "success": True,
            "format": "html",
            "content": html_content,
            "filename": f"test_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
            "size": len(html_content.encode('utf-8')),
            "content_type": "text/html"
        }
    
    def _export_tests_csv(self, tests: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Export tests as CSV"""
        csv_content = self._generate_tests_csv(tests)
        
        return {
            "success": True,
            "format": "csv",
            "content": csv_content,
            "filename": f"api_tests_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            "size": len(csv_content.encode('utf-8')),
            "content_type": "text/csv"
        }
    
    def _export_results_csv(self, results: Dict[str, Any]) -> Dict[str, Any]:
        """Export results as CSV"""
        csv_content = self._generate_results_csv(results)
        
        return {
            "success": True,
            "format": "csv",
            "content": csv_content,
            "filename": f"test_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            "size": len(csv_content.encode('utf-8')),
            "content_type": "text/csv"
        }
    
    def _export_fallback_text(self, data: List[Dict[str, Any]], export_type: str) -> Dict[str, Any]:
        """Fallback text export when other formats fail"""
        text_content = f"API {export_type.title()} Export\n"
        text_content += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        text_content += "=" * 50 + "\n\n"
        
        text_content += json.dumps(data, indent=2, default=str)
        
        return {
            "success": True,
            "format": "text",
            "content": text_content,
            "filename": f"api_{export_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "size": len(text_content.encode('utf-8')),
            "content_type": "text/plain"
        }
    
    def _generate_tests_html(self, tests: List[Dict[str, Any]]) -> str:
        """Generate HTML content for tests"""
        html = """<!DOCTYPE html>
<html>
<head>
    <title>API Test Suite</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 20px; }
        .header { text-align: center; margin-bottom: 30px; }
        .test { border: 1px solid #ddd; margin: 10px 0; padding: 15px; border-radius: 5px; }
        .test-header { background: #f5f5f5; padding: 10px; margin: -15px -15px 15px -15px; }
        .method { padding: 3px 8px; border-radius: 3px; color: white; font-weight: bold; }
        .method.GET { background: #61affe; }
        .method.POST { background: #49cc90; }
        .method.PUT { background: #fca130; }
        .method.DELETE { background: #f93e3e; }
        pre { background: #f8f8f8; padding: 10px; border-radius: 3px; overflow-x: auto; }
        .metadata { color: #666; font-size: 0.9em; }
    </style>
</head>
<body>
    <div class="header">
        <h1>API Test Suite</h1>
        <div class="metadata">
            Generated: """ + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + f"""<br>
            Total Tests: {len(tests)}
        </div>
    </div>
"""
        
        for i, test in enumerate(tests, 1):
            method = test.get('method', 'GET')
            html += f"""
    <div class="test">
        <div class="test-header">
            <h3>Test {i}: {test.get('name', 'Unnamed Test')}</h3>
            <span class="method {method}">{method}</span> {test.get('url', '')}
        </div>
        <p><strong>Description:</strong> {test.get('description', 'No description')}</p>
        <p><strong>Headers:</strong></p>
        <pre>{json.dumps(test.get('headers', {}), indent=2)}</pre>
"""
            
            if test.get('body'):
                html += f"""
        <p><strong>Request Body:</strong></p>
        <pre>{json.dumps(test.get('body'), indent=2)}</pre>
"""
            
            assertions = test.get('assertions', [])
            if assertions:
                html += f"""
        <p><strong>Assertions ({len(assertions)}):</strong></p>
        <ul>
"""
                for assertion in assertions:
                    html += f"            <li>{assertion.get('type', 'unknown')}: {assertion.get('description', 'No description')}</li>\n"
                
                html += "        </ul>\n"
            
            html += "    </div>\n"
        
        html += """
</body>
</html>"""
        
        return html
    
    def _generate_results_html(self, results: Dict[str, Any]) -> str:
        """Generate HTML content for results"""
        summary = results.get('summary', {})
        success_rate = summary.get('success_rate', 0)
        
        html = f"""<!DOCTYPE html>
<html>
<head>
    <title>API Test Results</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        .header {{ text-align: center; margin-bottom: 30px; }}
        .summary {{ background: #f5f5f5; padding: 20px; border-radius: 5px; margin-bottom: 20px; }}
        .metric {{ display: inline-block; margin: 10px 20px; text-align: center; }}
        .metric-value {{ font-size: 2em; font-weight: bold; }}
        .passed {{ color: #28a745; }}
        .failed {{ color: #dc3545; }}
        .test-result {{ border: 1px solid #ddd; margin: 10px 0; padding: 15px; border-radius: 5px; }}
        .test-result.passed {{ border-left: 5px solid #28a745; }}
        .test-result.failed {{ border-left: 5px solid #dc3545; }}
        .status {{ font-weight: bold; padding: 3px 8px; border-radius: 3px; color: white; }}
        .status.passed {{ background: #28a745; }}
        .status.failed {{ background: #dc3545; }}
        pre {{ background: #f8f8f8; padding: 10px; border-radius: 3px; overflow-x: auto; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>API Test Results Report</h1>
        <div>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
    </div>
    
    <div class="summary">
        <h2>Executive Summary</h2>
        <div class="metric">
            <div class="metric-value">{summary.get('total_tests', 0)}</div>
            <div>Total Tests</div>
        </div>
        <div class="metric">
            <div class="metric-value passed">{summary.get('passed', 0)}</div>
            <div>Passed</div>
        </div>
        <div class="metric">
            <div class="metric-value failed">{summary.get('failed', 0)}</div>
            <div>Failed</div>
        </div>
        <div class="metric">
            <div class="metric-value">{success_rate}%</div>
            <div>Success Rate</div>
        </div>
    </div>
"""
        
        test_results = results.get('test_results', [])
        if test_results:
            html += "    <h2>Detailed Results</h2>\n"
            
            for i, test_result in enumerate(test_results, 1):
                test = test_result.get('test', {})
                success = test_result.get('success', False)
                status_class = "passed" if success else "failed"
                status_text = "PASSED" if success else "FAILED"
                
                html += f"""
    <div class="test-result {status_class}">
        <h3>Test {i}: {test.get('name', 'Unnamed')} <span class="status {status_class}">{status_text}</span></h3>
        <p><strong>URL:</strong> {test.get('method', 'GET')} {test.get('url', '')}</p>
        <p><strong>Response Time:</strong> {test_result.get('execution_time', 0):.3f}s</p>
        <p><strong>Status Code:</strong> {test_result.get('response', {}).get('status_code', 'N/A')}</p>
"""
                
                if not success and test_result.get('error_message'):
                    html += f"        <p><strong>Error:</strong> <span class='failed'>{test_result.get('error_message')}</span></p>\n"
                
                html += "    </div>\n"
        
        html += """
</body>
</html>"""
        
        return html
    
    def _generate_tests_csv(self, tests: List[Dict[str, Any]]) -> str:
        """Generate CSV content for tests"""
        csv_lines = [
            "Name,Method,URL,Description,Headers,Body,Assertions,Tags,Timeout,Created"
        ]
        
        for test in tests:
            name = test.get('name', '').replace('"', '""')
            method = test.get('method', 'GET')
            url = test.get('url', '').replace('"', '""')
            description = test.get('description', '').replace('"', '""')
            headers = json.dumps(test.get('headers', {})).replace('"', '""')
            body = json.dumps(test.get('body')) if test.get('body') else ''
            body = body.replace('"', '""')
            assertions = str(len(test.get('assertions', [])))
            tags = ','.join(test.get('tags', []))
            timeout = str(test.get('timeout', 30))
            created = test.get('created_at', '')
            
            csv_lines.append(f'"{name}","{method}","{url}","{description}","{headers}","{body}","{assertions}","{tags}","{timeout}","{created}"')
        
        return '\n'.join(csv_lines)
    
    def _generate_results_csv(self, results: Dict[str, Any]) -> str:
        """Generate CSV content for results"""
        csv_lines = [
            "Test Name,Method,URL,Status,Response Time,Status Code,Assertions Passed,Assertions Failed,Error Message"
        ]
        
        test_results = results.get('test_results', [])
        for test_result in test_results:
            test = test_result.get('test', {})
            response = test_result.get('response', {})
            
            name = test.get('name', '').replace('"', '""')
            method = test.get('method', 'GET')
            url = test.get('url', '').replace('"', '""')
            status = "PASSED" if test_result.get('success', False) else "FAILED"
            response_time = f"{test_result.get('execution_time', 0):.3f}"
            status_code = str(response.get('status_code', ''))
            assertions_passed = str(test_result.get('assertions_passed', 0))
            assertions_failed = str(test_result.get('assertions_failed', 0))
            error_message = test_result.get('error_message', '').replace('"', '""')
            
            csv_lines.append(f'"{name}","{method}","{url}","{status}","{response_time}","{status_code}","{assertions_passed}","{assertions_failed}","{error_message}"')
        
        return '\n'.join(csv_lines)
