#!/usr/bin/env python3

"""
Enhanced Agent API Testing Platform Backend
Fixes issues with Postman collection import and test naming
"""

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import uuid
import json
import traceback
from datetime import datetime

# Create FastAPI app
app = FastAPI(title="Agent API Testing Platform - Enhanced", version="2.0.0")

# Enable CORS with explicit configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for development
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"]
)

# In-memory storage
stored_tests = {}

@app.get("/")
async def root():
    """Root endpoint for basic connectivity test"""
    return {
        "message": "Agent API Testing Platform - Enhanced",
        "status": "running", 
        "version": "2.0.0",
        "endpoints": {
            "health": "/health",
            "status": "/status",
            "tests": "/tests",
            "docs": "/docs"
        }
    }

@app.get("/health")
async def health_check():
    """Enhanced health check endpoint with detailed status"""
    return {
        "status": "healthy",
        "message": "Enhanced Agent API Testing Platform is running",
        "version": "2.0.0",
        "features": [
            "Postman Collection Import",
            "Batch Test Execution", 
            "AI-Powered Reports",
            "Performance Analytics"
        ],
        "endpoints": {
            "tests": len(stored_tests),
            "collections": len(set(test.get("collection", "default") for test in stored_tests.values())),
            "last_execution": "Ready"
        },
        "timestamp": datetime.now().isoformat()
    }

@app.get("/status")
async def get_status():
    """Backend status check"""
    return {
        "backend_status": "running",
        "platform": "Agent API Testing Platform Enhanced",
        "version": "2.0.0",
        "uptime": "Ready",
        "stored_tests": len(stored_tests),
        "available_endpoints": [
            "/health",
            "/tests", 
            "/import/tests",
            "/execute-test/{id}",
            "/execute-all-tests",
            "/execute-collection/{name}",
            "/generate-ai-report"
        ]
    }

@app.get("/tests")
async def get_tests():
    """Get all stored tests"""
    return {
        "status": "success",
        "tests": list(stored_tests.values()),
        "total": len(stored_tests)
    }

@app.post("/import/tests")
async def import_tests(request: dict):
    """Enhanced import with proper Postman collection support"""
    try:
        import_data = request.get("tests", {})
        
        print(f"\n=== ENHANCED IMPORT DEBUG ===")
        print(f"Import data type: {type(import_data)}")
        
        if not import_data:
            raise HTTPException(status_code=400, detail="No test data provided")
        
        imported_tests = []
        errors = []
        
        # Determine format and extract tests
        if isinstance(import_data, dict) and "info" in import_data and "item" in import_data:
            # Postman collection
            print(f"📦 Processing Postman Collection: {import_data['info'].get('name', 'Unknown')}")
            tests_to_import = extract_postman_tests(import_data)
        elif isinstance(import_data, dict) and "tests" in import_data:
            # Export format
            print("📁 Processing export format")
            tests_to_import = import_data["tests"]
        elif isinstance(import_data, list):
            # Direct array
            print(f"📋 Processing test array: {len(import_data)} items")
            tests_to_import = import_data
        else:
            # Single test
            print("📝 Processing single test")
            tests_to_import = [import_data]
        
        print(f"Total tests to process: {len(tests_to_import)}")
        
        # Process each test
        test_number = 0
        for index, test_data in enumerate(tests_to_import):
            try:
                test_number = index + 1
                print(f"\n--- Processing Test {test_number} ---")
                
                # Extract test information
                if "request" in test_data:
                    # Postman format
                    test_info = process_postman_test(test_data, test_number)
                else:
                    # Direct format
                    test_info = process_direct_test(test_data, test_number)
                
                if not test_info:
                    errors.append(f"Test {test_number}: Failed to process")
                    continue
                
                # Validate
                if not test_info.get("endpoint"):
                    errors.append(f"Test {test_number} ({test_info.get('name', 'Unknown')}): Missing endpoint")
                    continue
                
                # Create test object
                test_id = str(uuid.uuid4())
                current_time = datetime.now().isoformat()
                
                test_object = {
                    "id": test_id,
                    "name": test_info["name"],
                    "method": test_info["method"].upper(),
                    "endpoint": test_info["endpoint"],
                    "description": test_info.get("description", f"Imported {test_info['method']} test"),
                    "headers": test_info.get("headers", {}),
                    "body": test_info.get("body"),
                    "assertions": [{
                        "type": "status_code",
                        "expected": "200",
                        "description": "Status code should be 200"
                    }],
                    "tags": test_info.get("tags", ["imported"]),
                    "collection": test_info.get("collection", "imported-tests"),
                    "created_at": current_time,
                    "updated_at": current_time,
                    "imported_at": current_time
                }
                
                # Store
                stored_tests[test_id] = test_object
                imported_tests.append(test_object)
                
                print(f"✅ Imported: {test_object['name']} - {test_object['method']} {test_object['endpoint']}")
                
            except Exception as e:
                error_msg = f"Test {test_number}: {str(e)}"
                errors.append(error_msg)
                print(f"❌ Error: {error_msg}")
                traceback.print_exc()
        
        print(f"\n=== IMPORT COMPLETE ===")
        print(f"Successfully imported: {len(imported_tests)}")
        print(f"Errors: {len(errors)}")
        
        return {
            "status": "success",
            "imported_tests": len(imported_tests),
            "total_tests": len(imported_tests),
            "errors": errors,
            "tests": imported_tests[:3],
            "message": f"Successfully imported {len(imported_tests)} test(s)"
        }
        
    except Exception as e:
        print(f"❌ Import failed: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Import failed: {str(e)}")

def extract_postman_tests(collection):
    """Extract tests from Postman collection"""
    tests = []
    collection_name = collection.get("info", {}).get("name", "postman-collection")
    
    def process_items(items, folder_path=""):
        for item in items:
            if "item" in item:
                # Folder
                folder_name = item.get("name", "folder")
                new_path = f"{folder_path}/{folder_name}" if folder_path else folder_name
                process_items(item["item"], new_path)
            elif "request" in item:
                # Request
                test = {
                    "name": item.get("name", "Unnamed Request"),
                    "request": item["request"],
                    "collection": folder_path or collection_name
                }
                tests.append(test)
                print(f"Found request: {test['name']}")
    
    process_items(collection.get("item", []))
    return tests

def process_postman_test(test_data, test_number):
    """Process a Postman test request"""
    try:
        request = test_data["request"]
        
        # Extract name
        name = test_data.get("name", f"Postman Request {test_number}")
        
        # Extract method
        method = request.get("method", "GET")
        
        # Extract URL
        url_data = request.get("url", "")
        endpoint = extract_postman_url(url_data)
        
        # Extract headers
        headers = {}
        for header in request.get("header", []):
            if isinstance(header, dict) and not header.get("disabled", False):
                key = header.get("key", "")
                value = header.get("value", "")
                if key:
                    headers[key] = value
        
        # Extract body
        body = None
        body_data = request.get("body", {})
        if isinstance(body_data, dict):
            mode = body_data.get("mode", "")
            if mode == "raw":
                raw_data = body_data.get("raw", "")
                try:
                    body = json.loads(raw_data)
                except:
                    body = raw_data
            elif mode == "formdata":
                body = {}
                for item in body_data.get("formdata", []):
                    if isinstance(item, dict):
                        key = item.get("key")
                        value = item.get("value")
                        if key:
                            body[key] = value
        
        print(f"Postman test: {name} | {method} | {endpoint}")
        
        return {
            "name": name,
            "method": method,
            "endpoint": endpoint,
            "description": f"Imported from Postman: {name}",
            "headers": headers,
            "body": body,
            "tags": ["postman", "imported"],
            "collection": test_data.get("collection", "postman-import")
        }
        
    except Exception as e:
        print(f"Error processing Postman test: {e}")
        return None

def process_direct_test(test_data, test_number):
    """Process a direct test format"""
    try:
        name = (
            test_data.get("name") or 
            test_data.get("title") or 
            f"Test {test_number}"
        )
        
        method = (
            test_data.get("method") or 
            test_data.get("http_method") or 
            "GET"
        )
        
        endpoint = (
            test_data.get("endpoint") or 
            test_data.get("url") or 
            test_data.get("uri") or
            ""
        )
        
        print(f"Direct test: {name} | {method} | {endpoint}")
        
        return {
            "name": name,
            "method": method,
            "endpoint": endpoint,
            "description": test_data.get("description", f"Imported test: {name}"),
            "headers": test_data.get("headers", {}),
            "body": test_data.get("body"),
            "tags": test_data.get("tags", ["imported"]),
            "collection": test_data.get("collection", "imported-tests")
        }
        
    except Exception as e:
        print(f"Error processing direct test: {e}")
        return None

def extract_postman_url(url_data):
    """Extract URL from Postman URL format"""
    try:
        if isinstance(url_data, str):
            return url_data.strip()
        
        if isinstance(url_data, dict):
            # Try raw first
            if url_data.get("raw"):
                return url_data["raw"].strip()
            
            # Build from components
            protocol = url_data.get("protocol", "https")
            host = url_data.get("host", [])
            
            if isinstance(host, list):
                host_str = ".".join(str(h) for h in host if h)
            else:
                host_str = str(host)
            
            path = url_data.get("path", [])
            if isinstance(path, list):
                path_str = "/" + "/".join(str(p) for p in path if p)
            else:
                path_str = str(path) if path else ""
            
            # Add query parameters
            query = url_data.get("query", [])
            query_str = ""
            if isinstance(query, list) and query:
                params = []
                for q in query:
                    if isinstance(q, dict):
                        key = q.get("key", "")
                        value = q.get("value", "")
                        if key:
                            params.append(f"{key}={value}")
                if params:
                    query_str = "?" + "&".join(params)
            
            url = f"{protocol}://{host_str}{path_str}{query_str}"
            print(f"Constructed URL: {url}")
            return url
        
        return ""
        
    except Exception as e:
        print(f"Error extracting URL: {e}")
        return ""

@app.get("/tests/{test_id}")
async def get_test(test_id: str):
    """Get specific test"""
    if test_id not in stored_tests:
        raise HTTPException(status_code=404, detail="Test not found")
    return {"status": "success", "test": stored_tests[test_id]}

@app.delete("/tests/{test_id}")
async def delete_test(test_id: str):
    """Delete test"""
    if test_id not in stored_tests:
        raise HTTPException(status_code=404, detail="Test not found")
    
    deleted = stored_tests.pop(test_id)
    return {"status": "success", "message": f"Deleted test: {deleted['name']}"}

@app.post("/execute-test/{test_id}")
async def execute_test(test_id: str):
    """Execute a specific test"""
    if test_id not in stored_tests:
        raise HTTPException(status_code=404, detail="Test not found")
    
    test = stored_tests[test_id]
    
    try:
        import requests
        import time
        
        start_time = time.time()
        
        # Prepare request
        method = test["method"]
        url = test["endpoint"]
        headers = test.get("headers", {})
        body = test.get("body")
        
        print(f"🚀 Executing test: {test['name']}")
        print(f"   Method: {method}")
        print(f"   URL: {url}")
        print(f"   Headers: {headers}")
        print(f"   Body: {body}")
        
        # Make request
        response = requests.request(
            method=method,
            url=url,
            headers=headers,
            json=body if body else None,
            timeout=30
        )
        
        execution_time = time.time() - start_time
        
        # Process assertions
        assertion_results = []
        for assertion in test.get("assertions", []):
            assertion_type = ""
            expected = ""
            try:
                assertion_type = assertion.get("type", "")
                expected = assertion.get("expected", "")
                
                if assertion_type == "status_code":
                    actual = str(response.status_code)
                    passed = actual == str(expected)
                elif assertion_type == "response_time":
                    actual = execution_time * 1000  # Convert to ms
                    passed = actual <= float(expected)
                else:
                    passed = True
                    actual = "N/A"
                
                assertion_results.append({
                    "type": assertion_type,
                    "expected": expected,
                    "actual": actual,
                    "passed": passed,
                    "description": assertion.get("description", "")
                })
            except Exception as e:
                assertion_results.append({
                    "type": assertion_type,
                    "expected": expected,
                    "actual": f"Error: {str(e)}",
                    "passed": False,
                    "description": assertion.get("description", "")
                })
        
        # Overall success
        all_passed = all(a["passed"] for a in assertion_results)
        
        result = {
            "status": "success",
            "test": {
                "id": test_id,
                "name": test["name"],
                "method": test["method"],
                "endpoint": test["endpoint"]
            },
            "execution": {
                "success": all_passed,
                "status_code": response.status_code,
                "execution_time": round(execution_time, 3),
                "response_size": len(response.content),
                "timestamp": datetime.now().isoformat()
            },
            "assertions": assertion_results,
            "response": {
                "status_code": response.status_code,
                "headers": dict(response.headers),
                "body": response.text[:1000] if response.text else None  # First 1000 chars
            }
        }
        
        print(f"✅ Test execution completed: {all_passed}")
        return result
        
    except Exception as e:
        print(f"❌ Test execution failed: {e}")
        return {
            "status": "error",
            "test": {
                "id": test_id,
                "name": test["name"],
                "method": test["method"],
                "endpoint": test["endpoint"]
            },
            "execution": {
                "success": False,
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            },
            "message": f"Test execution failed: {str(e)}"
        }

@app.post("/create-sample-tests") 
async def create_sample_tests():
    """Create sample tests"""
    samples = [
        {
            "name": "JSONPlaceholder - Get Posts",
            "method": "GET",
            "endpoint": "https://jsonplaceholder.typicode.com/posts",
            "headers": {"Accept": "application/json"}
        },
        {
            "name": "JSONPlaceholder - Create Post", 
            "method": "POST",
            "endpoint": "https://jsonplaceholder.typicode.com/posts",
            "headers": {"Content-Type": "application/json"},
            "body": {"title": "Test", "body": "Content", "userId": 1}
        }
    ]
    
    created = []
    for sample in samples:
        test_id = str(uuid.uuid4())
        test_obj = {
            "id": test_id,
            "name": sample["name"],
            "method": sample["method"],
            "endpoint": sample["endpoint"],
            "description": f"Sample test: {sample['name']}",
            "headers": sample.get("headers", {}),
            "body": sample.get("body"),
            "assertions": [{"type": "status_code", "expected": "200", "description": "OK"}],
            "tags": ["sample"],
            "collection": "samples",
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat()
        }
        stored_tests[test_id] = test_obj
        created.append(test_obj)
    
    return {
        "status": "success",
        "created_tests": created,
        "message": f"Created {len(created)} sample tests"
    }

@app.put("/tests/{test_id}")
async def update_test(test_id: str, test_data: dict):
    """Update an existing test"""
    if test_id not in stored_tests:
        raise HTTPException(status_code=404, detail="Test not found")
    
    # Update the test
    stored_tests[test_id].update({
        "name": test_data.get("name", stored_tests[test_id]["name"]),
        "method": test_data.get("method", stored_tests[test_id]["method"]).upper(),
        "endpoint": test_data.get("endpoint", stored_tests[test_id]["endpoint"]),
        "description": test_data.get("description", stored_tests[test_id]["description"]),
        "headers": test_data.get("headers", stored_tests[test_id]["headers"]),
        "body": test_data.get("body", stored_tests[test_id]["body"]),
        "tags": test_data.get("tags", stored_tests[test_id]["tags"]),
        "collection": test_data.get("collection", stored_tests[test_id]["collection"]),
        "updated_at": datetime.now().isoformat()
    })
    
    return {
        "status": "success",
        "message": f"Test '{stored_tests[test_id]['name']}' updated successfully",
        "test": stored_tests[test_id]
    }

@app.post("/export/tests")
async def export_tests():
    """Export all tests"""
    export_data = {
        "export_info": {
            "exported_at": datetime.now().isoformat(),
            "total_tests": len(stored_tests),
            "platform": "Agent API Testing Platform"
        },
        "tests": list(stored_tests.values())
    }
    
    return {
        "status": "success",
        "export_data": export_data,
        "download_filename": f"api_tests_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    }

@app.post("/execute-collection/{collection_name}")
async def execute_collection(collection_name: str):
    """Execute all tests in a collection"""
    # Find tests in the collection
    collection_tests = [test for test in stored_tests.values() 
                       if test.get("collection") == collection_name]
    
    if not collection_tests:
        raise HTTPException(status_code=404, detail=f"No tests found in collection '{collection_name}'")
    
    print(f"🚀 Executing collection '{collection_name}' with {len(collection_tests)} tests")
    
    results = []
    for test in collection_tests:
        try:
            # Execute individual test
            execution_result = await execute_single_test(test)
            results.append(execution_result)
        except Exception as e:
            results.append({
                "test": {"id": test["id"], "name": test["name"]},
                "execution": {"success": False, "error": str(e)}
            })
    
    # Generate summary
    total_tests = len(results)
    passed_tests = sum(1 for r in results if r.get("execution", {}).get("success", False))
    failed_tests = total_tests - passed_tests
    
    return {
        "status": "success",
        "collection": collection_name,
        "summary": {
            "total_tests": total_tests,
            "passed": passed_tests,
            "failed": failed_tests,
            "success_rate": round((passed_tests / total_tests) * 100, 2) if total_tests > 0 else 0
        },
        "results": results,
        "executed_at": datetime.now().isoformat()
    }

@app.post("/execute-all-tests")
async def execute_all_tests():
    """Execute all stored tests"""
    if not stored_tests:
        raise HTTPException(status_code=404, detail="No tests available to execute")
    
    print(f"🚀 Executing all {len(stored_tests)} tests")
    
    results = []
    for test in stored_tests.values():
        try:
            execution_result = await execute_single_test(test)
            results.append(execution_result)
        except Exception as e:
            results.append({
                "test": {"id": test["id"], "name": test["name"]},
                "execution": {"success": False, "error": str(e)}
            })
    
    # Generate summary
    total_tests = len(results)
    passed_tests = sum(1 for r in results if r.get("execution", {}).get("success", False))
    failed_tests = total_tests - passed_tests
    
    return {
        "status": "success",
        "summary": {
            "total_tests": total_tests,
            "passed": passed_tests,
            "failed": failed_tests,
            "success_rate": round((passed_tests / total_tests) * 100, 2) if total_tests > 0 else 0
        },
        "results": results,
        "executed_at": datetime.now().isoformat()
    }

async def execute_single_test(test: dict):
    """Execute a single test and return results"""
    try:
        import requests
        import time
        
        start_time = time.time()
        
        # Make request
        response = requests.request(
            method=test["method"],
            url=test["endpoint"],
            headers=test.get("headers", {}),
            json=test.get("body") if test.get("body") else None,
            timeout=30
        )
        
        execution_time = time.time() - start_time
        
        # Process assertions
        assertion_results = []
        for assertion in test.get("assertions", []):
            assertion_type = assertion.get("type", "")
            expected = assertion.get("expected", "")
            
            if assertion_type == "status_code":
                actual = str(response.status_code)
                passed = actual == str(expected)
            elif assertion_type == "response_time":
                actual = execution_time * 1000
                passed = actual <= float(expected)
            else:
                passed = True
                actual = "N/A"
            
            assertion_results.append({
                "type": assertion_type,
                "expected": expected,
                "actual": actual,
                "passed": passed
            })
        
        all_passed = all(a["passed"] for a in assertion_results)
        
        return {
            "test": {
                "id": test["id"],
                "name": test["name"],
                "method": test["method"],
                "endpoint": test["endpoint"]
            },
            "execution": {
                "success": all_passed,
                "status_code": response.status_code,
                "execution_time": round(execution_time, 3),
                "response_size": len(response.content)
            },
            "assertions": assertion_results,
            "response": {
                "status_code": response.status_code,
                "headers": dict(response.headers),
                "body": response.text[:500] if response.text else None
            }
        }
        
    except Exception as e:
        return {
            "test": {
                "id": test["id"],
                "name": test["name"],
                "method": test["method"],
                "endpoint": test["endpoint"]
            },
            "execution": {
                "success": False,
                "error": str(e)
            }
        }

@app.post("/generate-ai-report")
async def generate_ai_report(request: dict):
    """Generate AI-powered Word report with test results and recommendations"""
    try:
        execution_results = request.get("execution_results", [])
        collection_name = request.get("collection_name", "API Tests")
        
        if not execution_results:
            raise HTTPException(status_code=400, detail="No execution results provided")
        
        # Generate AI suggestions
        ai_suggestions = generate_ai_suggestions(execution_results)
        
        # Create Word document
        report_data = create_word_report(execution_results, ai_suggestions, collection_name)
        
        return {
            "status": "success",
            "report": report_data,
            "download_url": f"/download-report/{report_data['report_id']}",
            "message": "AI-powered report generated successfully"
        }
        
    except Exception as e:
        print(f"❌ Report generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")

def generate_ai_suggestions(execution_results):
    """Generate AI suggestions based on test results"""
    suggestions = {
        "next_api_calls": [],
        "improvements": [],
        "security_recommendations": [],
        "performance_insights": []
    }
    
    # Analyze results and generate suggestions
    for result in execution_results:
        test = result.get("test", {})
        execution = result.get("execution", {})
        
        endpoint = test.get("endpoint", "")
        method = test.get("method", "")
        success = execution.get("success", False)
        status_code = execution.get("status_code", 0)
        response_time = execution.get("execution_time", 0)
        
        # Next API call suggestions
        if success and method == "POST":
            suggestions["next_api_calls"].append({
                "suggestion": f"Follow up with GET request to verify created resource",
                "endpoint": endpoint.replace("POST", "GET") if "POST" in endpoint else endpoint,
                "method": "GET",
                "reasoning": "Verify that the POST request successfully created the resource"
            })
        
        if success and method == "GET" and "/users" in endpoint:
            suggestions["next_api_calls"].append({
                "suggestion": "Test user-specific operations",
                "endpoint": endpoint + "/profile" if not endpoint.endswith("/") else endpoint + "profile",
                "method": "GET",
                "reasoning": "Test related user profile endpoints"
            })
        
        # Performance insights
        if response_time > 2.0:
            suggestions["performance_insights"].append({
                "issue": f"Slow response time: {response_time:.2f}s",
                "endpoint": endpoint,
                "recommendation": "Consider implementing caching or optimizing database queries"
            })
        
        # Security recommendations
        if status_code == 200 and method in ["PUT", "DELETE", "POST"]:
            suggestions["security_recommendations"].append({
                "endpoint": endpoint,
                "recommendation": "Ensure proper authentication and authorization for write operations",
                "priority": "High"
            })
    
    # General improvements
    failed_tests = [r for r in execution_results if not r.get("execution", {}).get("success", False)]
    if failed_tests:
        suggestions["improvements"].append({
            "issue": f"{len(failed_tests)} tests failed",
            "recommendation": "Review failed tests and implement proper error handling",
            "priority": "High"
        })
    
    return suggestions

def create_word_report(execution_results, ai_suggestions, collection_name):
    """Create a Word document report"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading(f'API Test Execution Report: {collection_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading('📊 Executive Summary', level=1)
        
        total_tests = len(execution_results)
        passed_tests = sum(1 for r in execution_results if r.get("execution", {}).get("success", False))
        failed_tests = total_tests - passed_tests
        success_rate = (passed_tests / total_tests) * 100 if total_tests > 0 else 0
        
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ("Total Tests", str(total_tests)),
            ("Passed Tests", str(passed_tests)),
            ("Failed Tests", str(failed_tests)),
            ("Success Rate", f"{success_rate:.1f}%"),
            ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
        
        # Test Results Details
        doc.add_heading('🧪 Test Results Details', level=1)
        
        for i, result in enumerate(execution_results):
            test = result.get("test", {})
            execution = result.get("execution", {})
            
            doc.add_heading(f'{i+1}. {test.get("name", "Unknown Test")}', level=2)
            
            details_table = doc.add_table(rows=6, cols=2)
            details_table.style = 'Table Grid'
            
            details_data = [
                ("Method", test.get("method", "N/A")),
                ("Endpoint", test.get("endpoint", "N/A")),
                ("Status", "✅ PASSED" if execution.get("success") else "❌ FAILED"),
                ("Response Code", str(execution.get("status_code", "N/A"))),
                ("Response Time", f"{execution.get('execution_time', 0):.3f}s"),
                ("Response Size", f"{execution.get('response_size', 0)} bytes")
            ]
            
            for j, (label, value) in enumerate(details_data):
                details_table.cell(j, 0).text = label
                details_table.cell(j, 1).text = value
        
        # AI Suggestions Section
        doc.add_heading('🤖 AI-Powered Recommendations', level=1)
        
        # Next API Calls
        if ai_suggestions["next_api_calls"]:
            doc.add_heading('🔮 Suggested Next API Calls', level=2)
            for suggestion in ai_suggestions["next_api_calls"]:
                p = doc.add_paragraph()
                p.add_run(f"• {suggestion['suggestion']}").bold = True
                doc.add_paragraph(f"   Method: {suggestion['method']}")
                doc.add_paragraph(f"   Endpoint: {suggestion['endpoint']}")
                doc.add_paragraph(f"   Reasoning: {suggestion['reasoning']}")
                doc.add_paragraph()
        
        # Performance Insights
        if ai_suggestions["performance_insights"]:
            doc.add_heading('⚡ Performance Insights', level=2)
            for insight in ai_suggestions["performance_insights"]:
                p = doc.add_paragraph()
                p.add_run(f"• {insight['issue']}").bold = True
                doc.add_paragraph(f"   Endpoint: {insight['endpoint']}")
                doc.add_paragraph(f"   Recommendation: {insight['recommendation']}")
                doc.add_paragraph()
        
        # Security Recommendations
        if ai_suggestions["security_recommendations"]:
            doc.add_heading('🔒 Security Recommendations', level=2)
            for rec in ai_suggestions["security_recommendations"]:
                p = doc.add_paragraph()
                p.add_run(f"• {rec['recommendation']}").bold = True
                doc.add_paragraph(f"   Endpoint: {rec['endpoint']}")
                doc.add_paragraph(f"   Priority: {rec['priority']}")
                doc.add_paragraph()
        
        # Save document
        report_id = str(uuid.uuid4())
        filename = f"api_test_report_{collection_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = f"/tmp/{filename}"
        
        doc.save(filepath)
        
        return {
            "report_id": report_id,
            "filename": filename,
            "filepath": filepath,
            "generated_at": datetime.now().isoformat(),
            "total_pages": len(doc.paragraphs) // 20,  # Rough estimate
            "file_size": "Estimated 2-5 MB"
        }
        
    except ImportError:
        # Fallback if python-docx is not installed
        return create_html_report(execution_results, ai_suggestions, collection_name)

def create_html_report(execution_results, ai_suggestions, collection_name):
    """Fallback HTML report if Word generation fails"""
    report_id = str(uuid.uuid4())
    
    total_tests = len(execution_results)
    passed_tests = sum(1 for r in execution_results if r.get("execution", {}).get("success", False))
    success_rate = (passed_tests / total_tests) * 100 if total_tests > 0 else 0
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>API Test Report: {collection_name}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; }}
            .header {{ text-align: center; color: #333; }}
            .summary {{ background: #f5f5f5; padding: 15px; border-radius: 5px; }}
            .test-result {{ margin: 15px 0; padding: 10px; border-left: 4px solid #ddd; }}
            .passed {{ border-left-color: #28a745; }}
            .failed {{ border-left-color: #dc3545; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>📊 API Test Execution Report</h1>
            <h2>{collection_name}</h2>
        </div>
        
        <div class="summary">
            <h3>📈 Summary</h3>
            <p><strong>Total Tests:</strong> {total_tests}</p>
            <p><strong>Passed:</strong> {passed_tests}</p>
            <p><strong>Failed:</strong> {total_tests - passed_tests}</p>
            <p><strong>Success Rate:</strong> {success_rate:.1f}%</p>
        </div>
        
        <h3>🧪 Test Results</h3>
    """
    
    for i, result in enumerate(execution_results):
        test = result.get("test", {})
        execution = result.get("execution", {})
        status_class = "passed" if execution.get("success") else "failed"
        
        html_content += f"""
        <div class="test-result {status_class}">
            <h4>{i+1}. {test.get('name', 'Unknown Test')}</h4>
            <p><strong>Method:</strong> {test.get('method')}</p>
            <p><strong>Endpoint:</strong> {test.get('endpoint')}</p>
            <p><strong>Status:</strong> {'✅ PASSED' if execution.get('success') else '❌ FAILED'}</p>
            <p><strong>Response Time:</strong> {execution.get('execution_time', 0):.3f}s</p>
        </div>
        """
    
    html_content += "</body></html>"
    
    # Save HTML file
    filename = f"api_test_report_{collection_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
    filepath = f"/tmp/{filename}"
    
    with open(filepath, 'w') as f:
        f.write(html_content)
    
    return {
        "report_id": report_id,
        "filename": filename,
        "filepath": filepath,
        "generated_at": datetime.now().isoformat(),
        "format": "HTML",
        "note": "Install python-docx for Word format: pip install python-docx"
    }

if __name__ == "__main__":
    import uvicorn
    print("🚀 Starting Enhanced Agent API Testing Platform...")
    print("📡 Backend: http://localhost:8000")
    print("📚 Docs: http://localhost:8000/docs")
    uvicorn.run("main_enhanced_import:app", host="0.0.0.0", port=8000, reload=True)